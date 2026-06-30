import subprocess
import sys
import os

def install(package):
    print(f"Installing {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install("python-docx")
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def format_run(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, space_after=12):
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=12)
    run = p.add_run(text)
    format_run(run, font_name="Times New Roman", size_pt=14, bold=True, color_rgb=(0,0,0))
    p.paragraph_format.keep_with_next = True
    return p

def main():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        
        # Configure Header
        header = sec.header
        hp = header.paragraphs[0]
        format_paragraph(hp, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
        hrun = hp.add_run("[TapTap (Blackbucks) Logo / Project Documentation]")
        format_run(hrun, font_name="Times New Roman", size_pt=9, italic=True, color_rgb=(100,100,100))
        
        # Configure Footer
        footer = sec.footer
        fp = footer.paragraphs[0]
        format_paragraph(fp, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        frun = fp.add_run("Page ")
        format_run(frun, font_name="Times New Roman", size_pt=10, color_rgb=(100,100,100))
        add_page_number(frun)

    # 1. COVER PAGE
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    run = p.add_run("\n\nAN INTERNSHIP PROJECT REPORT ON\n")
    format_run(run, size_pt=14, bold=True)
    
    run = p.add_run("AI-BASED JOB PORTAL\n\n")
    format_run(run, size_pt=18, bold=True, color_rgb=(79, 70, 229))
    
    run = p.add_run("Submitted in partial fulfillment of the requirements for the award of the degree of\n")
    format_run(run, size_pt=12, italic=True)
    
    run = p.add_run("BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE AND ENGINEERING\n\n")
    format_run(run, size_pt=14, bold=True)
    
    run = p.add_run("Submitted By\n")
    format_run(run, size_pt=12, bold=True)
    
    run = p.add_run("POLI NANDINI\n(Roll No: 23701A05D0)\n\n")
    format_run(run, size_pt=14, bold=True)
    
    run = p.add_run("Under the Guidance of\n")
    format_run(run, size_pt=12, bold=True)
    
    run = p.add_run("Sneha\n(Technical Guide)\n\n")
    format_run(run, size_pt=14, bold=True)
    
    run = p.add_run("Internship Organization: TapTap (Blackbucks)\n\n\n")
    format_run(run, size_pt=12, italic=True)
    
    run = p.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\n")
    format_run(run, size_pt=14, bold=True)
    
    run = p.add_run("ANNAMACHARYA INSTITUTE OF TECHNOLOGY AND SCIENCES (AUTONOMOUS)\n")
    format_run(run, size_pt=12, bold=True)
    
    run = p.add_run("Affiliated to JNTUA, Ananthapuramu\nRajampet - 516126, Andhra Pradesh\n")
    format_run(run, size_pt=12)
    
    run = p.add_run("Submission Date: 30 June 2026\n")
    format_run(run, size_pt=12, bold=True)
    
    doc.add_page_break()

    # 2. CERTIFICATE
    add_heading_styled(doc, "CERTIFICATE")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("This is to certify that the internship project report entitled \"AI-Based Job Portal\" is a bonafide record of work carried out by POLI NANDINI (Roll No: 23701A05D0) during their IV-I semester in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering from Annamacharya Institute of Technology and Sciences, Rajampet, affiliated to JNTUA, Ananthapuramu.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("\n\n\n\n\nSignature of Technical Guide                                                 Signature of Head of Department\n(Sneha)                                                                     (CSE Department)")
    format_run(run, bold=True)
    doc.add_page_break()

    # 3. ACKNOWLEDGEMENT
    add_heading_styled(doc, "ACKNOWLEDGEMENT")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("I express my deep gratitude to our guide Sneha for her valuable suggestions and supervision throughout the development of the AI-Based Job Portal project. I also thank Annamacharya Institute of Technology and Sciences for providing the academic support needed to complete my Java Full Stack Development internship.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("I extend my sincere thanks to TapTap (Blackbucks) for hosting my internship program and helping me build professional experience in software engineering, database design, and REST APIs integration.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("\n\nPOLI NANDINI\n(23701A05D0)")
    format_run(run, bold=True)
    doc.add_page_break()

    # 4. ABSTRACT
    add_heading_styled(doc, "ABSTRACT")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The AI-Based Job Portal is a full-stack web application designed to link candidates with employers based on their skill portfolios. Built using HTML, CSS, JavaScript on the frontend, and Java, Spring Boot, Spring Data JPA, and H2 on the backend, the system implements secure user authentication, resume file management, and semantic recommendation matching algorithms. Password security is achieved using BCrypt hashing, and the backend deploys a global exception advisor and validator framework to ensure robust data validation and clean handling. This report documents the system analysis, technical architecture, database designs, API specifications, and testing results of the portal.")
    format_run(run)
    doc.add_page_break()

    # 5. TABLE OF CONTENTS
    add_heading_styled(doc, "TABLE OF CONTENTS")
    p = doc.add_paragraph()
    format_paragraph(p)
    toc_lines = [
        ("1. Cover Page", "1"),
        ("2. Certificate", "2"),
        ("3. Acknowledgement", "3"),
        ("4. Abstract", "4"),
        ("5. Company Profile (TapTap / Blackbucks)", "6"),
        ("6. Internship Overview", "7"),
        ("7. Introduction & Problem Statement", "8"),
        ("8. System Requirements & Technology Stack", "10"),
        ("9. System Architecture & MVC design", "12"),
        ("10. Database Schema & UML Design", "14"),
        ("11. Frontend Modules Implementation", "17"),
        ("12. Backend Service & Controller Modules", "19"),
        ("13. REST API Specifications", "21"),
        ("14. Security & Cryptography (BCrypt)", "23"),
        ("15. Resume Upload & Download Logic", "24"),
        ("16. AI Matching Recommendation Service", "25"),
        ("17. Verification Testing & Case Matrix", "26"),
        ("18. Challenges Faced & Solutions", "28"),
        ("19. Conclusion & Future Enhancements", "29"),
        ("20. References", "30")
    ]
    for title, page in toc_lines:
        dots = "." * (80 - len(title) - len(page))
        line_p = doc.add_paragraph()
        format_paragraph(line_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
        run = line_p.add_run(f"{title} {dots} {page}")
        format_run(run)
    doc.add_page_break()

    # 6. COMPANY PROFILE
    add_heading_styled(doc, "COMPANY PROFILE (TAPTAP / BLACKBUCKS)")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("TapTap, an educational subsidiary of Blackbucks, is a leading software training and talent management organization focused on bridging academic curricula with real-world technology competencies. The organization specializes in FSD (Full Stack Development), machine learning implementations, cloud integrations, and mobile engineering training programs.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("During my internship at TapTap, I was mentored by Sneha, receiving hands-on guidance on industry standards, REST design principles, relational database models, git versioning controls, and unit validation frameworks. The organization fosters a collaborative agile environment that helps student interns acquire real-world problem-solving skills.")
    format_run(run)
    doc.add_page_break()

    # 7. INTERNSHIP OVERVIEW
    add_heading_styled(doc, "INTERNSHIP OVERVIEW")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("My internship in Java Full Stack Development (Java FSD) spanned a duration where I worked on designing, debugging, and testing full-stack web architectures. The program was divided into key phases:")
    format_run(run)
    
    phases = [
        "Phase 1: Database Models & JPA mappings (User, Job, Recruiter, Application tables setup).",
        "Phase 2: Backend Services creation (implementing business logic and constructor injections).",
        "Phase 3: Controller APIs mapping (Auth, User, Job, Application, Stats REST routes).",
        "Phase 4: Frontend fetch bindings (connecting HTML forms to port 8082 endpoints).",
        "Phase 5: Input Validation, BCrypt password hashing, global exception handlers, and resume file upload streams.",
        "Phase 6: UI/UX polishing, custom Toast notifies, search/location filters, and dashboards widgets."
    ]
    for phase in phases:
        lp = doc.add_paragraph(style='List Bullet')
        format_paragraph(lp, space_after=6)
        lrun = lp.add_run(phase)
        format_run(lrun)
    doc.add_page_break()

    # 8. INTRODUCTION
    add_heading_styled(doc, "INTRODUCTION")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The job search market relies heavily on high-speed, automated portals that help candidates filter roles and help recruiters analyze applicant pools. Traditional setups rely on localStorage or static representations that fail to persist data across sessions or protect credentials in the database. The AI-Based Job Portal acts as a robust carrier, providing persistent storage, password encryption, validation layers, exception handling, and custom resume file streams.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("By implementing a structured Spring Boot 3.3.0 backend paired with H2 local files database and a premium CSS3 glassmorphic client-side interface, the application bridges the gap between seekers and employers, returning matching analytics in real-time.")
    format_run(run)
    doc.add_page_break()

    # 9. PROBLEM STATEMENT & OBJECTIVES
    add_heading_styled(doc, "PROBLEM STATEMENT & OBJECTIVES")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Problem Statement: Legacy portals suffer from plain-text password vulnerabilities, memory leakages on local storage cache, lack of input filters (meaning users are flooded with irrelevant postings), lack of duplicate application controls, and high page load times due to unoptimized database operations.")
    format_run(run, bold=True)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Objectives:\n"
                     "1. Secure account operations using salt BCrypt encoders.\n"
                     "2. Set validation constraints at the model boundary to block empty values.\n"
                     "3. Establish a file repository uploader saving files to physical drives and download streams with custom attachment headers.\n"
                     "4. Formulate an automated data seeder ensuring dynamic startup entries.\n"
                     "5. Create a modern glassmorphic theme supporting responsive screens, toast warnings, and AJAX busy indicators.")
    format_run(run)
    doc.add_page_break()

    # 10. SYSTEM REQUIREMENTS & TECH STACK
    add_heading_styled(doc, "SYSTEM REQUIREMENTS & TECHNOLOGY STACK")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The development and execution requirements of the AI-Based Job Portal are detailed below:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Hardware Requirements:\n"
                     "- Processor: Intel Core i5/i7 (8th Gen+) or AMD Ryzen 5+\n"
                     "- Memory: 8 GB RAM minimum (16 GB recommended)\n"
                     "- Storage: 20 GB free Hard Disk/SSD space\n"
                     "- Network: Active port bindings capability (Tomcat binding to port 8082, Frontend on port 8081)")
    format_run(run)

    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Software Requirements:\n"
                     "- Operating System: Windows 10/11 or macOS/Linux\n"
                     "- Language SDK: Java Development Kit (JDK 17 or higher)\n"
                     "- Build Tool: Apache Maven 3.8+\n"
                     "- Database: File-based local H2 Engine\n"
                     "- Browsers: Google Chrome, Mozilla Firefox, Safari, Microsoft Edge")
    format_run(run)
    doc.add_page_break()

    # 11. SYSTEM ARCHITECTURE & MVC
    add_heading_styled(doc, "SYSTEM ARCHITECTURE & MVC DESIGN")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The project utilizes the Model-View-Controller (MVC) architectural design pattern:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("1. Model Layer: Handles the state and validation of database objects (User, Job, Recruiter, Application tables mapping).\n"
                     "2. View Layer: Displays pages dynamically using responsive glassmorphism CSS (HTML elements, toast alerts, loading spinners).\n"
                     "3. Controller Layer: Receives AJAX JSON and Multipart requests, calls service handlers, wraps data into DTOs, and returns HTTP responses.")
    format_run(run)

    # ASCII SYSTEM ARCHITECTURE
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(
        "┌────────────────────────────────────────────────────────┐\n"
        "│                   Frontend (HTML/CSS/JS)               │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ HTTP API REST Calls\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│                 Spring REST Controllers                │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ Service Bindings\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│                      Service Layer                     │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ Spring Data JPA\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│                     Repository Layer                   │\n"
        "└───────────────────────────┬────────────────────────────┘\n"
        "                            │ JDBC File Link\n"
        "                            ▼\n"
        "┌────────────────────────────────────────────────────────┐\n"
        "│                 H2 Persistent Database                 │\n"
        "└────────────────────────────────────────────────────────┘\n"
        "Figure 1: System Flow & MVC Architecture"
    )
    format_run(run, font_name="Courier New", size_pt=9.5, bold=True)
    doc.add_page_break()

    # 12. DATABASE DESIGN & ER DIAGRAM
    add_heading_styled(doc, "DATABASE DESIGN & ENTITY RELATIONSHIPS")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The database is structured to support relational integrity with H2's SQL compliance. Tables map User, Job, Recruiter, and Application profiles:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(
        "┌───────────────┐          1     *          ┌─────────────────┐\n"
        "│     USER      ├───────────────────────────┤   APPLICATION   │\n"
        "│ (Candidate)   │                           │                 │\n"
        "│               │                           │                 │\n"
        "│ - email (PK)  │                           │ - appId (PK)    │\n"
        "│ - name        │                           │ - userEmail (FK)│\n"
        "│ - password    │                           │ - jobTitle (FK) │\n"
        "│ - skills      │                           │ - status        │\n"
        "│ - resumeName  │                           └────────┬────────┘\n"
        "└───────────────┘                                    │ *\n"
        "                                                     │\n"
        "                                                     │ 1\n"
        "┌───────────────┐          1     *          ┌────────┴────────┐\n"
        "│   RECRUITER   ├───────────────────────────┤       JOB       │\n"
        "│ - email (PK)  │                           │ - id (PK)       │\n"
        "│ - companyName │                           │ - title         │\n"
        "│               │                           │ - company (FK)  │\n"
        "│               │                           │ - location      │\n"
        "└───────────────┘                           └─────────────────┘\n"
        "Figure 2: Entity Relationship Diagram (ERD)"
    )
    format_run(run, font_name="Courier New", size_pt=9.5, bold=True)
    doc.add_page_break()

    # 13. UML DIAGRAMS (USE CASE & ACTIVITY)
    add_heading_styled(doc, "UML DESIGN (USE CASE & ACTIVITY)")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Use Case Diagram describes the primary system actors (Candidate, Recruiter, Admin) and their respective actions:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(
        "           ┌───────────┐\n"
        "           │ Candidate ├─────────┐\n"
        "           └─────┬─────┘         │\n"
        "                 │               ▼\n"
        "                 ├────────► (Register / Login)\n"
        "                 ├────────► (Upload Resume)\n"
        "                 ├────────► (Search & Apply Jobs)\n"
        "                 └────────► (Get AI Matches)\n"
        "\n"
        "           ┌───────────┐\n"
        "           │ Recruiter ├─────────┐\n"
        "           └─────┬─────┘         ▼\n"
        "                 ├────────► (Post Job Position)\n"
        "                 └────────► (Monitor Metrics)\n"
        "\n"
        "           ┌───────────┐\n"
        "           │   Admin   ├─────────► (Manage Stats Table)\n"
        "           └───────────┘\n"
        "Figure 3: System Use Case Roles"
    )
    format_run(run, font_name="Courier New", size_pt=10, bold=True)
    doc.add_page_break()

    # 14. UML DIAGRAMS (SEQUENCE & CLASS)
    add_heading_styled(doc, "UML DESIGN (SEQUENCE & CLASS)")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Sequence Diagram tracks the registration and login validation execution flow:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(
        "Candidate              Browser              AuthController            AuthService            UserRepo\n"
        "   │                      │                       │                      │                     │\n"
        "   │─── Fill Credentials ─►                       │                      │                     │\n"
        "   │                      │─── POST /register ───►│                      │                     │\n"
        "   │                      │                       │─── check duplicate ─►│                     │\n"
        "   │                      │                       │                      │─── findByEmail ────►│\n"
        "   │                      │                       │                      │◄── returns null ────│\n"
        "   │                      │                       │◄── registerUser ─────│                     │\n"
        "   │                      │◄── Success Response ──│                      │                     │\n"
        "   │◄── Toast Success ────│                       │                      │                     │\n"
        "Figure 4: Registration Sequence Diagram"
    )
    format_run(run, font_name="Courier New", size_pt=9.5, bold=True)
    doc.add_page_break()

    # 15. FRONTEND MODULES IMPLEMENTATION
    add_heading_styled(doc, "FRONTEND MODULES IMPLEMENTATION")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The frontend directory consists of static HTML files paired with a centralized stylesheet (style.css) and action controller script (script.js). Elements utilize a cohesive layout:")
    format_run(run)
    
    modules = [
        "Navigation Header: Sticky, translucent glass navigation displaying user greeting welcome badges and logout triggers.",
        "Home Module: Statistics grids displaying platform application submissions count.",
        "Jobs Dashboard: Keyword search text inputs matched with company and location filter selectors.",
        "Profile Dashboard: Candidate name, email, skills display, dynamic resume uploader button, and download link.",
        "Toasts Module: Custom slide-in banner layout to visually notify users of API responses.",
        "Loaders Module: Full-screen modal blur with spinning widget overlays on network requests."
    ]
    for module in modules:
        lp = doc.add_paragraph(style='List Bullet')
        format_paragraph(lp, space_after=6)
        lrun = lp.add_run(module)
        format_run(lrun)
    doc.add_page_break()

    # 16. BACKEND SERVICE & CONTROLLER MODULES
    add_heading_styled(doc, "BACKEND SERVICE & CONTROLLER MODULES")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The backend application follows a decoupled tier pattern where controllers manage network requests while services handle business rules:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("1. Controllers: Bind DTO inputs, call service methods, and return status entities (e.g. AuthController, UserController, JobController, ApplicationController).\n"
                     "2. Services: Encapsulate transaction logic. Uses constructor injection (@Autowired avoided) to instantiate repository dependencies (e.g. UserService, AuthService, JobService, ApplicationService, RecommendationService).\n"
                     "3. Repositories: Interface layers extending JpaRepository to generate SQL queries (e.g. UserRepository, JobRepository, RecruiterRepository, ApplicationRepository).")
    format_run(run)
    doc.add_page_break()

    # 17. REST API SPECIFICATIONS
    add_heading_styled(doc, "REST API DOCUMENTATION")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The Spring Boot server exposes the following REST endpoint routes on Port 8082:")
    format_run(run)
    
    # Table of APIs
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Method", "Endpoint", "Request DTO", "Response DTO"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "4F46E5")
        for p_hdr in hdr_cells[i].paragraphs:
            format_paragraph(p_hdr, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            for r_hdr in p_hdr.runs:
                format_run(r_hdr, font_name="Times New Roman", size_pt=10, bold=True, color_rgb=(255,255,255))
                
    apis_data = [
        ("POST", "/register", "RegisterRequest", "String (Status)"),
        ("POST", "/login", "LoginRequest", "String (Status)"),
        ("GET", "/user/profile", "Query Param: email", "UserProfileResponse"),
        ("POST", "/user/profile/upload", "Multipart email + file", "UserProfileResponse"),
        ("POST", "/user/profile/update", "Query: email, name, skills", "UserProfileResponse"),
        ("GET", "/resume/{filename}", "Path Variable", "Attachment Stream"),
        ("GET", "/jobs", "None", "List<JobResponse>"),
        ("POST", "/jobs", "Job Entity", "Job Entity"),
        ("POST", "/applications", "Application Entity", "ApplicationResponse"),
        ("GET", "/recommendations", "Query Param: skills", "RecommendationResponse"),
        ("GET", "/admin/stats", "None", "Map<String, Object>")
    ]
    
    for method, endpoint, request_dto, response_dto in apis_data:
        row_cells = table.add_row().cells
        data_item = [method, endpoint, request_dto, response_dto]
        for idx, text in enumerate(data_item):
            row_cells[idx].text = text
            for p_cell in row_cells[idx].paragraphs:
                format_paragraph(p_cell, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
                for r_cell in p_cell.runs:
                    format_run(r_cell, font_name="Times New Roman", size_pt=9.5)
                    
    doc.add_paragraph() # Spacing
    doc.add_page_break()

    # 18. SECURITY & CRYPTOGRAPHY (BCRYPT)
    add_heading_styled(doc, "SECURITY & CRYPTOGRAPHY (BCRYPT)")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Protecting candidate credentials against leakage is a critical system objective. The project deploys BCryptPasswordEncoder from the Spring Security Crypto package. The password security process flows as follows:")
    format_run(run)
    
    steps = [
        "Registration Phase: The candidate submits a plain-text password. The AuthService hashes it using BCrypt with a secure random salt before storing it in H2 database.",
        "Login Validation Phase: The client submits their credentials. The server calls matches(rawPassword, encodedPassword) to verify the input password against the H2 database hash.",
        "Global Scope: Database leaks are fully mitigated as the passwords stored in H2 remain one-way encrypted."
    ]
    for step in steps:
        lp = doc.add_paragraph(style='List Bullet')
        format_paragraph(lp, space_after=6)
        lrun = lp.add_run(step)
        format_run(lrun)
    doc.add_page_break()

    # 19. RESUME UPLOAD & DOWNLOAD LOGIC
    add_heading_styled(doc, "RESUME UPLOAD & DOWNLOAD LOGIC")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The resume manager handles file operations in two main phases:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("File Upload: Sent via MultipartForm request. The backend checks if the file is empty, restricts size to 5 MB, and validates extensions (.pdf, .doc, .docx). Valid files are saved to the physical workspace folder `uploads/` using replace existing policy. The file name is updated in the candidate's database profile.")
    format_run(run)

    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("File Download: The controller resolves the path variable filename inside `uploads/`, verifies it exists and is readable, sets the content-type (e.g. application/pdf), and returns the file stream with header CONTENT_DISPOSITION as attachment, forcing a file download.")
    format_run(run)
    doc.add_page_break()

    # 20. AI MATCHING RECOMMENDATION SERVICE
    add_heading_styled(doc, "AI MATCHING RECOMMENDATION SERVICE")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The AI matching system uses a semantic ranking algorithm inside `RecommendationService`. It compares user skills against active postings. The ranking logic is:")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("- 85% Match: If the candidate's skills fully match the job title or company profile description.\n"
                     "- 50% Match: If there is a partial keyword overlap (e.g., candidate has 'Java' and job title includes 'Java Developer').\n"
                     "- 0% Match (No match found): If there are no common keywords.")
    format_run(run)
    
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("This ensures that candidates receive highly relevant job recommendations based on their skill sets, rather than random listings.")
    format_run(run)
    doc.add_page_break()

    # 21. VERIFICATION TESTING & CASE MATRIX
    add_heading_styled(doc, "VERIFICATION TESTING & CASE MATRIX")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Verification tests were executed across all system modules. The testing matrix is detailed below:")
    format_run(run)

    # Test cases table
    table_test = doc.add_table(rows=1, cols=5)
    table_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells_t = table_test.rows[0].cells
    hdr_titles_t = ["Test ID", "Module", "Expected Result", "Actual Result", "Status"]
    for i, title in enumerate(hdr_titles_t):
        hdr_cells_t[i].text = title
        set_cell_background(hdr_cells_t[i], "4F46E5")
        for p_hdr in hdr_cells_t[i].paragraphs:
            format_paragraph(p_hdr, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            for r_hdr in p_hdr.runs:
                format_run(r_hdr, font_name="Times New Roman", size_pt=10, bold=True, color_rgb=(255,255,255))
                
    tests_data = [
        ("TC-01", "Registration", "Saves encrypted password in database", "Saves BCrypt hash in User table", "PASS"),
        ("TC-02", "Login", "Validates credential match", "Permits correct hash match, rejects wrong", "PASS"),
        ("TC-03", "Validation", "Rejects email missing @ symbol", "Returns HTTP 400 with validation error", "PASS"),
        ("TC-04", "Duplicate App", "Blocks same email applying twice to a job", "Returns HTTP 400 'already applied'", "PASS"),
        ("TC-05", "Resume Upload", "Saves PDF in uploads/ folder", "File saved, database links path name", "PASS"),
        ("TC-06", "Resume Size", "Blocks file size exceeding 5 MB", "Returns HTTP 400 file limit exceeded", "PASS"),
        ("TC-07", "AI Matcher", "Calculates match score on skills", "Matches candidate skills against jobs", "PASS"),
        ("TC-08", "CORS Headers", "Permits cross origin fetches", "Allows port 8081 calling port 8082", "PASS")
    ]
    
    for tid, mod, exp, act, stat in tests_data:
        row_cells = table_test.add_row().cells
        data_item = [tid, mod, exp, act, stat]
        for idx, text in enumerate(data_item):
            row_cells[idx].text = text
            for p_cell in row_cells[idx].paragraphs:
                format_paragraph(p_cell, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
                for r_cell in p_cell.runs:
                    format_run(r_cell, font_name="Times New Roman", size_pt=9.5)
                    if idx == 4:
                         r_cell.bold = True
                         r_cell.font.color.rgb = RGBColor(16, 124, 65) # green
                         
    doc.add_paragraph() # Spacing
    doc.add_page_break()

    # 22. CHALLENGES FACED & SOLUTIONS
    add_heading_styled(doc, "CHALLENGES FACED & SOLUTIONS IMPLEMENTED")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("During development, several technical challenges were resolved:")
    format_run(run)
    
    challenges = [
        "CORS blocks on pre-controller exceptions: Resolved by creating a global WebMvcConfigurer bean system-wide rather than controller-level @CrossOrigin annotations, ensuring that all exceptions carry the required headers.",
        "Database seeder skips: Decoupled seeder checks to verify each table's record count individually rather than combined condition checks.",
        "File Size limitations: Increased default multipart size constraints in application.properties up to 10 MB.",
        "Invalid login feedback: Structured global exception handlers using @ControllerAdvice to return detailed ErrorDetails payloads."
    ]
    for ch in challenges:
        lp = doc.add_paragraph(style='List Bullet')
        format_paragraph(lp, space_after=6)
        lrun = lp.add_run(ch)
        format_run(lrun)
    doc.add_page_break()

    # 23. SKILLS LEARNED DURING INTERNSHIP
    add_heading_styled(doc, "SKILLS LEARNED DURING INTERNSHIP")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("The internship program at TapTap (Blackbucks) helped me build core technology competencies:")
    format_run(run)
    
    skills = [
        "Backend Architecture: Spring Boot FSD lifecycle, JPA repository patterns, dependency injection mechanisms, RESTful API mapping.",
        "Database Design: Schema mapping, data integrity, constraint checks, file-based database configurations.",
        "Frontend Engineering: DOM manipulation, asynchronous AJAX handling, CSS grid systems, responsive media queries.",
        "Security & Cryptography: BCrypt one-way password hashing.",
        "Development Operations: Git repository branching, Maven build scripts, application debugging."
    ]
    for sk in skills:
        lp = doc.add_paragraph(style='List Bullet')
        format_paragraph(lp, space_after=6)
        lrun = lp.add_run(sk)
        format_run(lrun)
    doc.add_page_break()

    # 24. CONCLUSION & FUTURE SCOPE
    add_heading_styled(doc, "CONCLUSION & FUTURE SCOPE")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Conclusion: The AI-Based Job Portal project is successfully completed and verified. The application supports multi-role access, database persistence, secure encryption, input validation constraints, and custom file streams. The glassmorphic interface is responsive and provides real-time search results, toast warnings, and loaders.")
    format_run(run)

    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("Future Scope: The portal can be extended with features such as JWT authentication, email verification links, PostgreSQL database migration, docker container deployment, and NLP-driven resume parsing.")
    format_run(run)
    doc.add_page_break()

    # 25. REFERENCES
    add_heading_styled(doc, "REFERENCES")
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run("1. Spring Boot Documentation: https://spring.io/projects/spring-boot\n"
                     "2. Hibernate ORM Core: https://hibernate.org/orm/\n"
                     "3. MDN Web Docs (JavaScript Guide): https://developer.mozilla.org\n"
                     "4. JNTUA Curriculum Guidelines for Internship Project Reports.\n"
                     "5. TapTap (Blackbucks) Full Stack Java Development Training Manual.")
    format_run(run)

    # Save to the root directory
    output_filename = "AI_Based_Job_Portal_Documentation.docx"
    output_path = os.path.join("..", output_filename)
    
    # Try saving, if path resolution fails, save in current directory
    try:
        doc.save(output_path)
        print(f"SUCCESS: Document saved to {output_path}")
    except Exception as e:
        print(f"Warning: Could not save to {output_path}, saving in current folder instead: {e}")
        doc.save(output_filename)
        print(f"SUCCESS: Document saved to {output_filename}")

if __name__ == "__main__":
    main()
